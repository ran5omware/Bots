import sqlite3
from flask import Flask, render_template, redirect, request, flash, send_from_directory, url_for
from werkzeug.exceptions import abort
import os
import pprint
from docx import Document
from datetime import datetime

DIRECTORY = os.getcwd() + "\\"
NUMBER = 0

app = Flask(__name__)

TEMP_DICT = {
    "name": '',
    "passport": '',
    "pc_id": ''
}


# Функция для подключения к базе данных
def get_db_connection(db_name):
    conn = sqlite3.connect(db_name)
    conn.row_factory = sqlite3.Row
    return conn


# Главная страница
@app.route('/')
def index():
    return render_template("index.html")


@app.route('/contracts')
def contracts():
    conn = get_db_connection('database.db')
    inv = conn.execute('SELECT * FROM Contracts').fetchall()
    conn.close()
    return render_template("contracts.html", contract=inv)


@app.route('/contracts/<int:contract_id>')
def one_contract(contract_id):
    conn = get_db_connection('database.db')
    contract = conn.execute('SELECT * FROM Contracts WHERE id=?', (contract_id,)).fetchall()[0]
    print(contract['id'])
    conn.close()
    return render_template("one_contract.html", contract=contract)


@app.route('/contracts/create_contract', methods=['GET', 'POST'])
def create_contract():
    if request.method == 'POST':
        print(1)
        conn = get_db_connection('database.db')
        try:
            previous_employee = conn.execute('SELECT * FROM computers').fetchall()[-1]
            new_id = previous_employee['id'] + 1
        except IndexError:
            new_id = 0
        number = request.form.get('number')
        date = request.form['date']
        deal_type = request.form['deal_type']
        start_price = request.form['start_price']
        discount = request.form['discount']
        deal_status = request.form['deal_status']
        finall_price = request.form['finall_price']
        pc_id = request.form['id']
        client_id = request.form['client_id']
        employee_id = request.form['employee_id']
        cursor = conn.cursor()
        cursor.execute("INSERT INTO Contracts (id_contract, number, date, deal_type, start_price, discount, deal_status, finall_price, id, client_id, employee_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                       (new_id, number, date, deal_type, start_price, discount, deal_status, finall_price, pc_id,
                        client_id, employee_id))
        conn.commit()
        conn.close()
        return redirect(url_for('contracts'))
    return render_template('create_contract.html')


@app.route('/computers')
def computers():
    conn = get_db_connection('database.db')
    list_of_computers = conn.execute('SELECT * FROM Computers').fetchall()
    conn.close()
    return render_template("computers.html", computers=list_of_computers)


@app.route('/computers/<int:id>')
def one_computer(id):
    conn = get_db_connection('database.db')
    computer = conn.execute('SELECT * FROM Computers').fetchall()[id]
    conn.close()
    return render_template("one_computer.html", computer=computer)


@app.route('/computers/create_computer', methods=['GET', 'POST'])
def create_computer():
    if request.method == 'POST':
        conn = get_db_connection('database.db')
        try:
            previous_employee = conn.execute('SELECT * FROM Computers').fetchall()[-1]
            new_id = previous_employee['id'] + 1
        except IndexError:
            new_id = 0
        GPU = request.form['CPU']
        CPU = request.form['GPU']
        Motherboard = request.form['Motherboard']
        RAM = request.form['RAM']
        Case = request.form['Case']
        available = request.form['available']
        client_id = request.form['client_id']
        price = request.form['price']
        cursor = conn.cursor()
        cursor.execute("INSERT INTO Computers (id, CPU, GPU, Motherboard, RAM, \"Case\", available, client_id, price) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                       (new_id, CPU, GPU, Motherboard, RAM, Case, available, client_id, price))
        conn.commit()
        conn.close()
        return redirect(url_for('computers'))
    return render_template('create_computer.html')


# Функция удаления строк
@app.route('/remove/<table>-<int:row_id>', methods=['GET', 'POST'])
def remove_row(table, row_id=-1):
    conn = get_db_connection('database.db')
    cur = conn.cursor()
    cur.execute(f'DELETE FROM {table} WHERE id=?', (row_id,))
    conn.commit()
    conn.close()
    return redirect(url_for(table))


@app.route('/clients')
def clients():
    return render_template("clients.html")


@app.route('/employees')
def employees():
    conn = get_db_connection('database.db')
    empls = conn.execute('SELECT * FROM Employees').fetchall()
    conn.close()
    return render_template("employees.html", employees=empls)


@app.route('/employees/<int:employee_id>')
def one_employee(employee_id):
    conn = get_db_connection('database.db')
    employee = conn.execute('SELECT * FROM Employees').fetchall()[employee_id]
    conn.close()
    return render_template("one_employee.html", employee=employee)


@app.route('/employees/create_employee', methods=['GET', 'POST'])
def create_employee():
    if request.method == 'POST':
        conn = get_db_connection('database.db')
        try:
            previous_employee = conn.execute('SELECT * FROM Employees').fetchall()[-1]
            new_id = previous_employee['id'] + 1
        except IndexError:
            new_id = 0
        name = request.form.get('name')
        email = request.form['email']
        phone_number = request.form['phone_number']
        position = request.form['position']
        department = request.form['department']
        chief_id = request.form['chief_id']
        cursor = conn.cursor()
        cursor.execute('INSERT INTO Employees (id, name, email, phone_number, position, department, chief_id)'
                       ' VALUES (?, ?, ?, ?, ?, ?, ?)',
                       (new_id, name, email, phone_number, position, department, chief_id))
        conn.commit()
        conn.close()
        return redirect(url_for('employees'))
    return render_template('create_employee.html')


def get_data_for_report():
    conn = get_db_connection('database.db')
    computers = conn.execute('SELECT * FROM Computers WHERE id=?',(TEMP_DICT['pc_id'],)).fetchone()
    person = conn.execute('SELECT * FROM Employees WHERE department=?', ('Хозяйственный Отдел',)).fetchall()
    print(person)
    conn.close()
    id = computers['id']
    price = computers['price']
    client_id = computers['client_id']
    available = computers['available']
    return [id, price, client_id, available], [person[0]['name'], person[0]['position']]


@app.route('/create_report', methods=['GET', 'POST'])
def create_report():
    if request.method == 'POST':
        name = request.form.get('name')
        passport = request.form['passport']
        pc_id = request.form['id']
        TEMP_DICT['name'] = name
        TEMP_DICT['passport'] = passport
        TEMP_DICT['pc_id'] = pc_id
        return redirect(url_for('save_report'))
    return render_template("create_report.html")


@app.route('/save_report')
def save_report():
    number = 0
    table_data, sub_data = get_data_for_report()
    template_name = DIRECTORY + "docx_templates\\template_report.docx"
    template_doc = Document(template_name)
    data = {}
    with open("docx_templates/report_keys.txt", 'rt', encoding='utf-8') as data_txt:
        not_so_keys = [elem.strip() for elem in data_txt.readlines()]
        for k in range(len(not_so_keys)):
            data[not_so_keys[k]] = ""
    data["{{REPORT_NUMBER}}"] = str(number)
    data["{{DATE}}"] = " ".join(list(reversed(str(datetime.today().date()).split("-"))))
    # for i in range(len(table_data)):
    #     tab = template_doc.tables[0]
    #     insert_rows(tab, table_data[i], i + 1)
    table = template_doc.tables[0]
    table.add_row()
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = str(table_data[0])
    table.cell(1, 2).text = str(table_data[1])
    table.cell(1, 3).text = str(table_data[2])
    table.cell(1, 4).text = str(table_data[3])
    # data['{{NUMBER}}'] = "1"
    # data['{{PC_ID}}'] = table_data[0]
    # data['{{PRICE}}'] = table_data[1]
    # data['{{CLIENT_ID}}'] = table_data[2]
    # data['{{AVAILABLE}}'] = table_data[3]
    data["{{EMPLOYEE_NAME}}"] = sub_data[0]
    data["{{EMPLOYEE_POSITION}}"] = sub_data[1]
    for key, value in data.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, key, value)

    template_doc.save(DIRECTORY + "\\reports\\Договор о продаже №" + str(number) + ".docx")
    return redirect(url_for('index'))


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


# def insert_rows(table, table_data, num):
#     table.add_row()
#     print(table_data)
#     pc_id, price, client_id, available = table_data
#     table.cell(num, 0).text = str(num)
#     table.cell(num, 1).text = str(pc_id)
#     table.cell(num, 2).text = str(price)
#     table.cell(num, 3).text = str(client_id)
#     table.cell(num, 4).text = str(available)
#     # return price * volume


def remove_sub_script():
    try:
        os.remove(DIRECTORY + "\\Отчет о наличии №0.docx")
    except FileNotFoundError:
        pass


if __name__ == '__main__':
    remove_sub_script()
    print(DIRECTORY)
    app.run()
